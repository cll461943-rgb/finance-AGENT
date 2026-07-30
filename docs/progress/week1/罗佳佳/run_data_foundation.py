#!/usr/bin/env python3
"""Build corpus inventory, parsing PoC, evidence samples and quality logs."""

import argparse
import csv
import hashlib
import json
import re
import subprocess
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
import xlrd
from docx import Document
from jsonschema import Draft202012Validator
from pypdf import PdfReader

PARSER_VERSION = "1.0.0"
SUPPORTED = {".xls", ".xlsx", ".pdf", ".doc", ".docx"}
CN_MONTH = re.compile(r"(?<!\d)(1[0-2]|[1-9])月")
YEAR = re.compile(r"(20\d{2}|201\d)年")
QUARTER = re.compile(r"(?:第|[一二三四1-4])?([一二三四1-4])季度|([1-4])季")
CLAUSE = re.compile(r"^第[一二三四五六七八九十百零〇0-9]+条")


def sha256(path):
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def title_from_name(name):
    stem = Path(name).stem
    stem = re.sub(r"^\d{3}_", "", stem)
    parts = stem.split("_", 1)
    return parts[0].strip() if parts else stem


def category(name):
    labels = []
    rules = [
        ("保险业报表", r"保险业|人身险|财产险|原保险|偿付能力|资金运用"),
        ("银行业报表", r"银行业.*(?:资产|负债|指标|贷款)|商业银行主要"),
        ("监管制度", r"办法|通知|规定|指引|规则|细则|规程"),
        ("资本管理", r"资本管理|资本工具|风险加权"),
        ("发布日程", r"发布日程"),
        ("绿色信贷", r"绿色信贷"),
        ("年报", r"年报")
    ]
    for label, pattern in rules:
        if re.search(pattern, name):
            labels.append(label)
    return "|".join(labels) or "其他"


def time_fields(name):
    ym = YEAR.search(name)
    mm = CN_MONTH.search(name)
    qm = QUARTER.search(name)
    qmap = {"一": 1, "二": 2, "三": 3, "四": 4}
    qraw = next((x for x in qm.groups() if x), None) if qm else None
    return (
        int(ym.group(1)) if ym else None,
        int(mm.group(1)) if mm else None,
        qmap.get(qraw, int(qraw) if qraw and qraw.isdigit() else None)
    )


def basic_doc_id(path, digest):
    m = re.match(r"(\d{3})_", path.name)
    return f"DOC-{m.group(1) if m else '000'}-{digest[:12]}"


def excel_profile(path):
    profile = {"sheet_count": 0, "page_count": None, "table_count": 0,
               "merged_range_count": 0, "formula_count": 0, "nonempty_cell_count": 0,
               "has_multirow_header": False, "warnings": []}
    if path.suffix.lower() == ".xlsx":
        wb = openpyxl.load_workbook(path, read_only=False, data_only=False)
        profile["sheet_count"] = len(wb.worksheets)
        for ws in wb.worksheets:
            profile["merged_range_count"] += len(ws.merged_cells.ranges)
            profile["has_multirow_header"] |= any(r.max_row > r.min_row for r in ws.merged_cells.ranges)
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value not in (None, ""):
                        profile["nonempty_cell_count"] += 1
                    if cell.data_type == "f":
                        profile["formula_count"] += 1
        profile["table_count"] = profile["sheet_count"]
        wb.close()
    else:
        wb = xlrd.open_workbook(path, formatting_info=True, on_demand=True)
        profile["sheet_count"] = wb.nsheets
        for sname in wb.sheet_names():
            ws = wb.sheet_by_name(sname)
            profile["merged_range_count"] += len(ws.merged_cells)
            profile["has_multirow_header"] |= any(rhi - rlo > 1 for rlo, rhi, _, _ in ws.merged_cells)
            profile["nonempty_cell_count"] += sum(
                1 for r in range(ws.nrows) for c in range(ws.ncols)
                if ws.cell_value(r, c) not in (None, "")
            )
        profile["table_count"] = profile["sheet_count"]
        wb.release_resources()
    return profile


def pdf_profile(path):
    reader = PdfReader(str(path))
    encrypted = reader.is_encrypted
    if encrypted:
        try:
            reader.decrypt("")
        except Exception:
            pass
    lengths = []
    for page in reader.pages[:5]:
        lengths.append(len((page.extract_text() or "").strip()))
    return {"sheet_count": None, "page_count": len(reader.pages), "table_count": None,
            "merged_range_count": None, "formula_count": None,
            "nonempty_cell_count": None, "has_multirow_header": None,
            "warnings": (["encrypted_pdf"] if encrypted else []) +
                        (["possible_scanned_pdf"] if lengths and max(lengths) < 30 else [])}


def word_profile(path):
    if path.suffix.lower() == ".docx":
        doc = Document(str(path))
        para_count = sum(1 for p in doc.paragraphs if p.text.strip())
        return {"sheet_count": None, "page_count": None, "table_count": len(doc.tables),
                "merged_range_count": None, "formula_count": None,
                "nonempty_cell_count": para_count, "has_multirow_header": None, "warnings": []}
    result = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(path)],
                            capture_output=True, check=False)
    text = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0:
        raise RuntimeError(result.stderr.decode("utf-8", errors="replace")[:300])
    return {"sheet_count": None, "page_count": None, "table_count": None,
            "merged_range_count": None, "formula_count": None,
            "nonempty_cell_count": len([x for x in text.splitlines() if x.strip()]),
            "has_multirow_header": None,
            "warnings": ["legacy_doc_layout_not_preserved"]}


def profile(path):
    ext = path.suffix.lower()
    if ext in {".xls", ".xlsx"}:
        return excel_profile(path)
    if ext == ".pdf":
        return pdf_profile(path)
    return word_profile(path)


def inventory(corpus):
    records = []
    for path in sorted(p for p in corpus.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED):
        digest = sha256(path)
        y, m, q = time_fields(path.name)
        error = None
        try:
            details = profile(path)
            parseable = True
        except Exception as exc:
            details = {"sheet_count": None, "page_count": None, "table_count": None,
                       "merged_range_count": None, "formula_count": None,
                       "nonempty_cell_count": None, "has_multirow_header": None,
                       "warnings": []}
            parseable = False
            error = f"{type(exc).__name__}: {exc}"[:500]
        number = re.match(r"(\d{3})_", path.name)
        is_attachment = bool(re.search(r"附件|样表|报告模板|材料目录", path.name))
        records.append({
            "doc_id": basic_doc_id(path, digest), "file_no": number.group(1) if number else "",
            "filename": path.name, "title": title_from_name(path.name),
            # Store a portable corpus-relative path; never leak a developer's
            # machine-specific absolute workspace path into manifests/evidence.
            "local_path": str(Path(corpus.name) / path.name), "_source_path": str(path.resolve()),
            "extension": path.suffix.lower(),
            "size_bytes": path.stat().st_size, "sha256": digest, "category": category(path.name),
            "year": y, "month": m, "quarter": q, "is_attachment": is_attachment,
            "parent_doc_id": "", "source_url": "", "parseable": parseable,
            "parse_error": error or "", **details
        })
    return records


def pick_samples(records):
    chosen = []
    def add(rec, reason):
        if rec and rec["doc_id"] not in {x[0]["doc_id"] for x in chosen}:
            chosen.append((rec, reason))
    for ext in [".xls", ".xlsx", ".pdf", ".doc", ".docx"]:
        group = [r for r in records if r["extension"] == ext]
        if group:
            add(group[0], f"{ext}基础样本")
            add(max(group, key=lambda x: x["size_bytes"]), f"{ext}最大文件")
            add(group[len(group)//2], f"{ext}中位序号样本")
    excels = [r for r in records if r["extension"] in {".xls", ".xlsx"}]
    add(max(excels, key=lambda x: x["sheet_count"] or 0), "最多Sheet")
    add(max(excels, key=lambda x: x["merged_range_count"] or 0), "最多合并单元格")
    add(max(excels, key=lambda x: x["formula_count"] or 0), "最多公式")
    attachments = [r for r in records if r["is_attachment"]]
    if attachments:
        add(attachments[0], "附件关系样本")
    scanned = [r for r in records if "possible_scanned_pdf" in r["warnings"]]
    if scanned:
        add(scanned[0], "疑似扫描PDF")
    return chosen


def base_evidence(rec, eid, etype, content, location, table=None, warnings=None):
    return {
        "schema_version": "1.0.0", "evidence_id": eid, "doc_id": rec["doc_id"],
        "evidence_type": etype,
        "source": {"title": rec["title"], "local_path": rec["local_path"],
                   "source_url": None, "file_sha256": rec["sha256"]},
        "content": content,
        "location": {"page": None, "chapter": None, "section": None, "clause_no": None,
                     "char_start": None, "char_end": None, "sheet_name": None, "cell_range": None,
                     **location},
        "table_semantics": {"row_header": None, "column_header": None, "header_path": [],
                            "value": None, "unit": None, "period": None, "scale": None,
                            "footnote": None, "formula": None, **(table or {})},
        "regulation_metadata": {"agency": None, "publish_date": None, "effective_date": None,
                                "expire_date": None, "status": "unknown", "version_relation": None},
        "quality": {"validation_status": "auto_checked", "support_type": "unknown",
                    "sufficiency": "unknown", "conflict_group": None,
                    "parser_name": "regrag_data_foundation", "parser_version": PARSER_VERSION,
                    "warnings": warnings or []}
    }


def excel_evidence(rec, limit=8):
    path = Path(rec.get("_source_path", rec["local_path"]))
    rows = []
    if path.suffix.lower() == ".xlsx":
        wb = openpyxl.load_workbook(path, read_only=False, data_only=False)
        for ws in wb.worksheets:
            values = [[c.value for c in row] for row in ws.iter_rows()]
            for ri, row in enumerate(values):
                for ci, value in enumerate(row):
                    if value not in (None, "") and (isinstance(value, (int, float)) or (isinstance(value, str) and value.startswith("="))):
                        coord = openpyxl.utils.get_column_letter(ci + 1) + str(ri + 1)
                        rows.append((ws.title, coord, value, values, ri, ci))
                        if len(rows) >= limit: break
                if len(rows) >= limit: break
            if len(rows) >= limit: break
        wb.close()
    else:
        wb = xlrd.open_workbook(path, formatting_info=True, on_demand=True)
        for sname in wb.sheet_names():
            ws = wb.sheet_by_name(sname)
            values = [[ws.cell_value(r, c) for c in range(ws.ncols)] for r in range(ws.nrows)]
            for ri, row in enumerate(values):
                for ci, value in enumerate(row):
                    if isinstance(value, (int, float)) and value not in (0, ""):
                        coord = openpyxl.utils.get_column_letter(ci + 1) + str(ri + 1)
                        rows.append((sname, coord, value, values, ri, ci))
                        if len(rows) >= limit: break
                if len(rows) >= limit: break
            if len(rows) >= limit: break
        wb.release_resources()
    evidence = []
    for idx, (sheet, coord, value, values, ri, ci) in enumerate(rows, 1):
        # Headers must be textual labels. Numeric values above/left are neighboring
        # observations, not headers; treating them as headers corrupts Cell-RAG.
        row_header = next((str(values[ri][j]).strip() for j in range(ci - 1, -1, -1)
                           if isinstance(values[ri][j], str) and values[ri][j].strip()), None)
        col_headers = [str(values[r][ci]).strip() for r in range(0, ri)
                       if ci < len(values[r]) and isinstance(values[r][ci], str)
                       and values[r][ci].strip()]
        col_headers = col_headers[-3:]
        unit_match = re.search(r"单位[：:]?\s*([^\s，。；]+)", " ".join(str(x) for row in values[:8] for x in row if x not in (None, "")))
        period_parts = []
        if rec.get("year"):
            period_parts.append(str(rec["year"]))
        month_header = next((x for x in reversed(col_headers) if CN_MONTH.search(x)), None)
        if month_header:
            period_parts.append(f"{int(CN_MONTH.search(month_header).group(1)):02d}")
        table = {"row_header": row_header, "column_header": col_headers[-1] if col_headers else None,
                 "header_path": col_headers, "value": str(value),
                 "unit": unit_match.group(1) if unit_match else None,
                 "period": "-".join(period_parts) if period_parts else None,
                 "formula": value if isinstance(value, str) and value.startswith("=") else None}
        content = f"{row_header or ''} | {' / '.join(col_headers)} | {value}".strip(" |")
        evidence.append(base_evidence(rec, f"{rec['doc_id']}-CELL-{idx:03d}", "table_cell", content,
                                      {"sheet_name": sheet, "cell_range": coord}, table))
    return evidence


def text_evidence(rec, limit=8):
    path = Path(rec.get("_source_path", rec["local_path"]))
    chunks = []
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        for pno, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            for para in re.split(r"\n+", text):
                if len(para.strip()) >= 20:
                    chunks.append((para.strip(), {"page": pno}))
                    if len(chunks) >= limit: break
            if len(chunks) >= limit: break
    elif path.suffix.lower() == ".docx":
        doc = Document(str(path))
        offset = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                chunks.append((text, {"char_start": offset, "char_end": offset + len(text)}))
                offset += len(text) + 1
                if len(chunks) >= limit: break
    else:
        result = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(path)], capture_output=True, check=True)
        text = result.stdout.decode("utf-8", errors="replace")
        for match in re.finditer(r"[^\n]+", text):
            para = match.group().strip()
            if para:
                chunks.append((para, {"char_start": match.start(), "char_end": match.end()}))
                if len(chunks) >= limit: break
    evidence = []
    for idx, (text, location) in enumerate(chunks, 1):
        cm = CLAUSE.match(text)
        if cm:
            location["clause_no"] = cm.group()
        evidence.append(base_evidence(rec, f"{rec['doc_id']}-TEXT-{idx:03d}",
                                      "clause" if cm else "paragraph", text, location,
                                      warnings=["legacy_doc_layout_not_preserved"] if path.suffix.lower() == ".doc" else []))
    return evidence


def write_csv(path, rows, fields):
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader(); writer.writerows(rows)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--corpus", required=True, type=Path)
    ap.add_argument("--output", default=Path(__file__).parent / "outputs", type=Path)
    args = ap.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)
    records = inventory(args.corpus)
    fields = [field for field in records[0].keys() if not field.startswith("_")]
    for r in records:
        r["warnings"] = "|".join(r["warnings"])
    write_csv(args.output / "corpus_manifest.csv", records, fields)
    samples = pick_samples([{**r, "warnings": r["warnings"].split("|") if r["warnings"] else []} for r in records])
    evidence, quality, sample_rows = [], [], []
    run_id = datetime.now(timezone.utc).strftime("RUN-%Y%m%dT%H%M%SZ")
    for rec, reason in samples:
        started = datetime.now(timezone.utc)
        try:
            ev = excel_evidence(rec) if rec["extension"] in {".xls", ".xlsx"} else text_evidence(rec)
            status = "warning" if rec["warnings"] else ("passed" if ev else "warning")
            err_type = "SOURCE_WARNING" if rec["warnings"] else ("" if ev else "EMPTY_EVIDENCE")
            err_detail = "|".join(rec["warnings"]) if rec["warnings"] else ("" if ev else "解析成功但未生成非空证据")
        except Exception as exc:
            ev, status, err_type = [], "failed", type(exc).__name__
            err_detail = str(exc)[:500]
        evidence.extend(ev)
        elapsed = round((datetime.now(timezone.utc) - started).total_seconds() * 1000, 2)
        sample_rows.append({"doc_id": rec["doc_id"], "filename": rec["filename"], "extension": rec["extension"],
                            "selection_reason": reason, "parse_status": status, "evidence_count": len(ev),
                            "elapsed_ms": elapsed, "warnings": "|".join(rec["warnings"])})
        quality.append({"run_id": run_id, "parser_version": PARSER_VERSION, "doc_id": rec["doc_id"],
                        "filename": rec["filename"], "stage": "poc_parse", "parse_status": status,
                        "error_type": err_type, "error_detail": err_detail, "expected_value": "至少1条可定位证据",
                        "actual_value": f"{len(ev)}条", "affected_location": "", "severity": "high" if status == "failed" else "",
                        "review_method": "automatic", "reviewer": "", "fix_status": "review_required" if status == "warning" else ("open" if status == "failed" else "not_required"),
                        "fixed_version": "", "review_time": ""})
    # Corpus-level quality findings belong in the same auditable ledger.
    hashes = {}
    for rec in records:
        hashes.setdefault(rec["sha256"], []).append(rec)
    for digest, dupes in hashes.items():
        if len(dupes) > 1:
            names = " | ".join(x["filename"] for x in dupes)
            for rec in dupes:
                quality.append({"run_id": run_id, "parser_version": PARSER_VERSION, "doc_id": rec["doc_id"],
                                "filename": rec["filename"], "stage": "inventory", "parse_status": "warning",
                                "error_type": "DUPLICATE_CONTENT", "error_detail": names,
                                "expected_value": "文件内容唯一", "actual_value": digest,
                                "affected_location": "whole_file", "severity": "high", "review_method": "sha256",
                                "reviewer": "", "fix_status": "review_required", "fixed_version": "", "review_time": ""})
    missing_ids = sorted(set(range(1, 508)) - {int(r["file_no"]) for r in records if r["file_no"]})
    quality.append({"run_id": run_id, "parser_version": PARSER_VERSION, "doc_id": "CORPUS",
                    "filename": "", "stage": "inventory", "parse_status": "warning",
                    "error_type": "MISSING_SEQUENCE_IDS", "error_detail": ",".join(f"{x:03d}" for x in missing_ids),
                    "expected_value": "001-507连续", "actual_value": f"缺失{len(missing_ids)}个编号",
                    "affected_location": "corpus", "severity": "informational", "review_method": "automatic",
                    "reviewer": "", "fix_status": "accepted_source_condition", "fixed_version": "", "review_time": ""})
    unlinked = [r for r in records if r["is_attachment"] and not r["parent_doc_id"]]
    quality.append({"run_id": run_id, "parser_version": PARSER_VERSION, "doc_id": "CORPUS",
                    "filename": "", "stage": "inventory", "parse_status": "warning",
                    "error_type": "UNLINKED_ATTACHMENTS", "error_detail": "附件来源标题存在于文件名，但语料中通常无独立主文档，不能可靠生成parent_doc_id",
                    "expected_value": "附件关联主文档", "actual_value": f"{len(unlinked)}份未关联",
                    "affected_location": "corpus", "severity": "high", "review_method": "automatic",
                    "reviewer": "", "fix_status": "review_required", "fixed_version": "", "review_time": ""})
    odd_names = [r for r in records if re.search(r"\.x\.xlsx$", r["filename"], re.I)]
    for rec in odd_names:
        quality.append({"run_id": run_id, "parser_version": PARSER_VERSION, "doc_id": rec["doc_id"],
                        "filename": rec["filename"], "stage": "inventory", "parse_status": "warning",
                        "error_type": "SUSPICIOUS_FILENAME", "error_detail": "文件名以.x.xlsx结尾",
                        "expected_value": "规范扩展名", "actual_value": rec["filename"],
                        "affected_location": "filename", "severity": "low", "review_method": "automatic",
                        "reviewer": "", "fix_status": "review_required", "fixed_version": "", "review_time": ""})
    schema = json.loads((Path(__file__).parent / "evidence_schema.json").read_text(encoding="utf-8"))
    validator = Draft202012Validator(schema)
    schema_errors = []
    for ev in evidence:
        for error in validator.iter_errors(ev):
            schema_errors.append({"evidence_id": ev["evidence_id"], "path": "/".join(map(str, error.path)), "message": error.message})
    with (args.output / "evidence_samples.jsonl").open("w", encoding="utf-8") as f:
        for ev in evidence: f.write(json.dumps(ev, ensure_ascii=False) + "\n")
    write_csv(args.output / "poc_samples.csv", sample_rows, list(sample_rows[0].keys()))
    write_csv(args.output / "parse_quality_log.csv", quality, list(quality[0].keys()))
    write_csv(args.output / "schema_validation_errors.csv", schema_errors,
              ["evidence_id", "path", "message"])
    ext_counts = Counter(r["extension"] for r in records)
    missing = missing_ids
    passed = sum(1 for r in sample_rows if r["parse_status"] == "passed")
    report = f"""# 数据底座交付报告

生成时间：{datetime.now().astimezone().isoformat(timespec='seconds')}  
解析器版本：{PARSER_VERSION}

## 1. 数据清点

- 文件总数：{len(records)}
- 格式分布：{', '.join(f'{k}={v}' for k, v in sorted(ext_counts.items()))}
- 可打开并完成结构探测：{sum(bool(r['parseable']) for r in records)}/{len(records)}
- 编号范围：001–507；缺失编号：{', '.join(f'{x:03d}' for x in missing)}
- 重复 SHA-256 文件数：{sum(v - 1 for v in Counter(r['sha256'] for r in records).values() if v > 1)}

完整逐文件记录见 `corpus_manifest.csv`。

## 2. 解析 PoC

- 分层样本：{len(sample_rows)} 份，覆盖 xls/xlsx/pdf/doc/docx、最大文件、复杂 Excel 和附件。
- 通过：{passed}；需人工复核：{sum(r['parse_status']=='warning' for r in sample_rows)}；失败：{sum(r['parse_status']=='failed' for r in sample_rows)}。
- 生成证据：{len(evidence)} 条，保存在 `evidence_samples.jsonl`。
- Schema 校验错误：{len(schema_errors)} 条。

PoC 是自动结构与可定位性检查，不代替人工语义验收。下一步应在 `parse_quality_log.csv` 补充 reviewer、review_time，并重点核对数值、单位、期间、多级表头和条款边界。

## 3. 证据结构

`evidence_schema.json` 固定了身份、来源、位置、表格语义、制度版本及可信状态。
`evidence_samples.jsonl` 保存本次 PoC 从 {len(sample_rows)} 份代表性文件中生成的 {len(evidence)} 条真实证据样例，用于验证 Schema、测试后续检索入库和展示来源定位；它不是 500 份语料的全量证据库。
Excel 证据精确到 Sheet + Cell，PDF 精确到页，Word 精确到字符区间；旧版 DOC 由 textutil 转换，版面结构可能损失并已标警告。

## 4. 质量记录

`parse_quality_log.csv` 是本次 PoC 的逐文件质量台账；`schema_validation_errors.csv` 记录结构校验问题。`poc_samples.csv` 保存抽样理由、耗时和证据数量，便于复现。

## 5. 交付物清单

- 数据清点：`corpus_manifest.csv`
- 解析 PoC：`poc_samples.csv`、`run_data_foundation.py`
- 证据结构：`evidence_schema.json`、`evidence_samples.jsonl`
- 质量记录：`parse_quality_log.csv`、`schema_validation_errors.csv`
- 使用与复现：`README.md`、`requirements.txt`

## 验收结论

{'自动 PoC 已全部通过，可以进入人工抽检。' if passed == len(sample_rows) and not schema_errors else '结构解析可运行，但存在需人工复核或修复的质量项；进入全量入库前请处理质量日志中的 review_required/open 项。'}
"""
    (args.output / "交付报告.md").write_text(report, encoding="utf-8")
    print(json.dumps({"files": len(records), "samples": len(sample_rows), "passed": passed,
                      "evidence": len(evidence), "schema_errors": len(schema_errors),
                      "output": str(args.output.resolve())}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
